"""
Script to convert PRESENTATION_GUIDE.md to Word document
with RTL for Arabic and LTR for English
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def set_rtl(paragraph):
    """Set paragraph to RTL (Right-to-Left) for Arabic"""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '1')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def set_ltr(paragraph):
    """Set paragraph to LTR (Left-to-Right) for English"""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    bidi.set(qn('w:val'), '0')
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

def is_arabic(text):
    """Check if text contains Arabic characters"""
    arabic_pattern = re.compile(r'[\u0600-\u06FF\u0750-\u077F\u08A0-\u08FF]+')
    return bool(arabic_pattern.search(text))

def add_heading(doc, text, level):
    """Add heading with proper direction"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    run.font.size = Pt(18 - level * 2)
    run.font.bold = True

    if is_arabic(text):
        set_rtl(heading)
        run.font.name = 'Arial'
    else:
        set_ltr(heading)
        run.font.name = 'Consolas'

    return heading

def add_paragraph_with_direction(doc, text, is_quote=False, is_code=False):
    """Add paragraph with proper RTL/LTR direction"""
    para = doc.add_paragraph()

    if is_code:
        run = para.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        para.paragraph_format.left_indent = Inches(0.5)
        set_ltr(para)
        # Add gray background for code
        shading = OxmlElement('w:shd')
        shading.set(qn('w:fill'), 'F0F0F0')
        para._p.get_or_add_pPr().append(shading)
    elif is_quote:
        run = para.add_run(text)
        run.font.size = Pt(11)
        para.paragraph_format.left_indent = Inches(0.3)
        if is_arabic(text):
            set_rtl(para)
            run.font.name = 'Arial'
            para.paragraph_format.right_indent = Inches(0.3)
        else:
            set_ltr(para)
            run.font.name = 'Calibri'
    else:
        run = para.add_run(text)
        run.font.size = Pt(11)
        if is_arabic(text):
            set_rtl(para)
            run.font.name = 'Arial'
        else:
            set_ltr(para)
            run.font.name = 'Calibri'

    return para

def create_word_document():
    """Create Word document from presentation guide"""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # Title
    title = doc.add_heading(level=0)
    run = title.add_run('دليل العرض التقديمي الكامل')
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.name = 'Arial'
    set_rtl(title)

    subtitle = doc.add_paragraph()
    run = subtitle.add_run('نظام الحضور الذكي - Smart Attendance System')
    run.font.size = Pt(16)
    run.font.name = 'Arial'
    set_rtl(subtitle)

    doc.add_paragraph('─' * 50)

    # ═══════════════════════════════════════════════════════════════
    # Section: How to Start
    # ═══════════════════════════════════════════════════════════════

    add_heading(doc, 'كيف تبدأ العرض', 1)

    add_heading(doc, 'قبل العرض (تحضير):', 2)
    steps = [
        '1. شغّل RabbitMQ: rabbitmq-server (أو من Windows Services)',
        '2. شغّل النظام: .\\START.ps1',
        '3. تأكد من فتح: http://localhost:3000',
        '4. سجّل الدخول: admin / admin123',
        '5. افتح هذا الملف أمامك للقراءة'
    ]
    for step in steps:
        add_paragraph_with_direction(doc, step)

    add_heading(doc, 'المقدمة (قلها للمشرف):', 2)

    # Arabic intro
    p = add_paragraph_with_direction(doc,
        'بالعربي: "هذا نظام حضور ذكي مبني على معمارية الخدمات المصغرة. يتكون من 8 خدمات مستقلة، وطبّقت فيه 4 أنماط تصميم هي: Circuit Breaker باستخدام State Pattern، و Strategy Pattern مع Reflection، و Choreography Pattern باستخدام RabbitMQ، و JWT Authentication."',
        is_quote=True)

    # English intro
    p = add_paragraph_with_direction(doc,
        'In English: "This is a smart attendance system built on microservices architecture. It consists of 8 independent services, and I implemented 4 design patterns: Circuit Breaker using State Pattern, Strategy Pattern with Reflection, Choreography Pattern using RabbitMQ, and JWT Authentication."',
        is_quote=True)

    add_heading(doc, 'ترتيب الشرح المقترح:', 2)
    order = [
        '1. JWT Authentication (الأسهل - ابدأ به)',
        '2. Strategy Pattern (واضح ومباشر)',
        '3. Circuit Breaker (أكثر تعقيداً)',
        '4. Choreography (الأكثر تعقيداً - اتركه للنهاية)'
    ]
    for item in order:
        add_paragraph_with_direction(doc, item)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # Requirement 1: Circuit Breaker
    # ═══════════════════════════════════════════════════════════════

    add_heading(doc, '═' * 30, 1)
    add_heading(doc, 'المطلب الأول: Circuit Breaker Pattern (State Pattern)', 1)
    add_heading(doc, '═' * 30, 1)

    add_heading(doc, 'ماذا تقول للمشرف (المقدمة):', 2)
    add_paragraph_with_direction(doc,
        'بالعربي: "هذا النمط يحمي النظام من الانهيار المتتالي. مثلاً إذا خدمة الحضور تعطلت، بدل ما نضيّع وقت ننتظرها كل مرة، النظام يتذكر إنها معطلة ويرفض الطلبات فوراً لمدة معينة، ثم يعيد الاختبار."',
        is_quote=True)
    add_paragraph_with_direction(doc,
        'In English: "This pattern protects the system from cascading failures. For example, if the attendance service fails, instead of wasting time waiting for it every time, the system remembers it\'s down and rejects requests immediately for a certain period, then retests."',
        is_quote=True)

    add_heading(doc, 'ماذا تقول (الحالات الثلاث):', 2)
    add_paragraph_with_direction(doc,
        'بالعربي: "النمط له 3 حالات مثل State Pattern:',
        is_quote=True)
    add_paragraph_with_direction(doc, '  - CLOSED: الحالة الطبيعية، الطلبات تمر', is_quote=True)
    add_paragraph_with_direction(doc, '  - OPEN: بعد 3 أخطاء متتالية، الطلبات تُرفض فوراً', is_quote=True)
    add_paragraph_with_direction(doc, '  - HALF_OPEN: بعد 15 ثانية، نجرب طلب واحد لنرى هل الخدمة رجعت"', is_quote=True)

    add_paragraph_with_direction(doc,
        'In English: "The pattern has 3 states like State Pattern:',
        is_quote=True)
    add_paragraph_with_direction(doc, '  - CLOSED: Normal state, requests pass through', is_quote=True)
    add_paragraph_with_direction(doc, '  - OPEN: After 3 consecutive failures, requests are rejected immediately', is_quote=True)
    add_paragraph_with_direction(doc, '  - HALF_OPEN: After 15 seconds, we try one request to see if the service recovered"', is_quote=True)

    add_heading(doc, 'الملفات ومواقع الكود', 2)

    add_heading(doc, 'الملف الرئيسي (التطبيق اليدوي):', 3)
    add_paragraph_with_direction(doc,
        'المسار الكامل: c:\\Users\\HP\\smart-attendance-system\\common\\circuit_breaker.py')

    add_heading(doc, 'تعريف الحالات (سطر 19-24):', 3)
    code1 = '''class CircuitState(Enum):
    """Circuit breaker states"""
    CLOSED = "closed"          # Normal operation
    OPEN = "open"              # Failing - reject requests
    HALF_OPEN = "half_open"    # Testing if service recovered'''
    add_paragraph_with_direction(doc, code1, is_code=True)

    add_paragraph_with_direction(doc, 'قل للمشرف:', is_quote=False)
    add_paragraph_with_direction(doc, 'بالعربي: "هنا عرّفت الحالات الثلاث كـ Enum، وهذا أساس State Pattern."', is_quote=True)
    add_paragraph_with_direction(doc, 'In English: "Here I defined the three states as an Enum, and this is the foundation of State Pattern."', is_quote=True)

    add_heading(doc, 'الدالة الرئيسية call (سطر 77-103):', 3)
    code2 = '''def call(self, func: Callable, *args, **kwargs) -> Any:
    # Check if circuit is OPEN
    if self.state == CircuitState.OPEN:
        if self._should_attempt_reset():
            self.state = CircuitState.HALF_OPEN  # Test transition
        else:
            raise Exception("Circuit breaker is OPEN")  # Immediate reject

    # Try to execute the function
    try:
        result = func(*args, **kwargs)
        self._on_success()
        return result
    except Exception as e:
        self._on_failure()
        raise e'''
    add_paragraph_with_direction(doc, code2, is_code=True)

    add_paragraph_with_direction(doc, 'قل للمشرف:', is_quote=False)
    add_paragraph_with_direction(doc, 'بالعربي: "هذه الدالة الأساسية. أولاً تتحقق من الحالة، إذا OPEN ترفض فوراً، وإلا تحاول تنفيذ الطلب."', is_quote=True)
    add_paragraph_with_direction(doc, 'In English: "This is the main function. First it checks the state, if OPEN it rejects immediately, otherwise it tries to execute the request."', is_quote=True)

    add_heading(doc, 'دالة معالجة الفشل (سطر 123-139):', 3)
    code3 = '''def _on_failure(self):
    """Handle failed call"""
    self.failure_count += 1
    self.last_failure_time = datetime.now()

    if self.state == CircuitState.HALF_OPEN:
        self.state = CircuitState.OPEN  # Test failed

    elif self.failure_count >= self.failure_threshold:
        self.state = CircuitState.OPEN  # 3 errors = open circuit'''
    add_paragraph_with_direction(doc, code3, is_code=True)

    add_paragraph_with_direction(doc, 'قل للمشرف:', is_quote=False)
    add_paragraph_with_direction(doc, 'بالعربي: "عند كل فشل، نزيد العداد. إذا وصل 3، نفتح الدائرة."', is_quote=True)
    add_paragraph_with_direction(doc, 'In English: "On each failure, we increment the counter. If it reaches 3, we open the circuit."', is_quote=True)

    add_heading(doc, 'الملف الثاني (التطبيق بالمكتبة):', 3)
    add_paragraph_with_direction(doc,
        'المسار الكامل: c:\\Users\\HP\\smart-attendance-system\\common\\circuit_breaker_library.py')

    code4 = '''self.breaker = pybreaker.CircuitBreaker(
    fail_max=failure_threshold,    # 3 errors = open circuit
    reset_timeout=timeout,         # 15 seconds then test
    listeners=[self.listener],     # Logging listener
    name=name
)'''
    add_paragraph_with_direction(doc, code4, is_code=True)

    add_paragraph_with_direction(doc, 'قل للمشرف:', is_quote=False)
    add_paragraph_with_direction(doc, 'بالعربي: "هنا نفس المفهوم لكن باستخدام مكتبة pybreaker. أقل كود، نفس النتيجة."', is_quote=True)
    add_paragraph_with_direction(doc, 'In English: "Here\'s the same concept but using the pybreaker library. Less code, same result."', is_quote=True)

    add_heading(doc, 'كيف تُظهر للمشرف (خطوة بخطوة):', 2)

    demo_steps_cb = [
        ('الخطوة 1: افتح Dashboard',
         'قل: "سأفتح صفحة Design Patterns ثم Circuit Breaker"',
         'Say: "I\'ll open the Design Patterns page then Circuit Breaker"',
         'http://localhost:3000 → Design Patterns → Circuit Breaker tab'),
        ('الخطوة 2: أظهر الحالة الطبيعية',
         'قل: "هنا الحالة CLOSED، يعني كل شي طبيعي"',
         'Say: "Here the state is CLOSED, meaning everything is normal"',
         'أشر على البطاقة الخضراء'),
        ('الخطوة 3: أوقف Attendance Service',
         'قل: "الآن سأوقف خدمة الحضور لمحاكاة عطل"',
         'Say: "Now I\'ll stop the attendance service to simulate a failure"',
         'أغلق نافذة Attendance Service'),
        ('الخطوة 4: اضغط Test 3 مرات',
         'قل: "كل ضغطة = خطأ. بعد 3 أخطاء، الدائرة تنفتح"',
         'Say: "Each click = error. After 3 errors, the circuit opens"',
         'اضغط "Test Attendance Service" 3 مرات - أشر على تغيّر الحالة إلى OPEN (أحمر)'),
        ('الخطوة 5: أظهر الرفض الفوري',
         'قل: "الآن لاحظ: الطلبات تُرفض فوراً بدون انتظار"',
         'Say: "Now notice: requests are rejected immediately without waiting"',
         'اضغط Test مرة أخرى - أشر على الرسالة: "Circuit is OPEN"'),
        ('الخطوة 6: انتظر HALF_OPEN',
         'قل: "بعد 15 ثانية، النظام يجرب مرة أخرى"',
         'Say: "After 15 seconds, the system tries again"',
         'انتظر حتى تتحول إلى HALF_OPEN (أصفر)'),
        ('الخطوة 7: أعد تشغيل الخدمة',
         'قل: "الآن أشغّل الخدمة ونختبر"',
         'Say: "Now I\'ll start the service and test"',
         'شغّل Attendance Service - اضغط Test مرتين بنجاح - أشر على تغيّر الحالة إلى CLOSED (أخضر)')
    ]

    for step in demo_steps_cb:
        add_heading(doc, step[0], 3)
        add_paragraph_with_direction(doc, step[1], is_quote=True)
        add_paragraph_with_direction(doc, step[2], is_quote=True)
        add_paragraph_with_direction(doc, step[3])

    add_heading(doc, 'أسئلة متوقعة:', 2)

    qa_cb = [
        ('س: لماذا 3 أخطاء وليس 5؟',
         'بالعربي: "رقم متوازن. لو أقل، الدائرة تفتح بخطأ عابر. لو أكثر، نتأخر في اكتشاف المشكلة."',
         'In English: "It\'s a balanced number. If less, the circuit opens on a transient error. If more, we delay detecting the problem."'),
        ('س: لماذا 15 ثانية؟',
         'بالعربي: "يعطي الخدمة وقت للتعافي. لو أقل، نضغط على خدمة مريضة. لو أكثر، ننتظر كثير."',
         'In English: "It gives the service time to recover. If less, we pressure a sick service. If more, we wait too long."'),
        ('س: ما الفرق بين التطبيق اليدوي والمكتبة؟',
         'بالعربي: "اليدوي: تحكم كامل وتعلم أفضل. المكتبة: أسرع وأقل أخطاء."',
         'In English: "Manual: full control and better learning. Library: faster and fewer bugs."'),
        ('س: هذا Client-side أم Server-side؟',
         'بالعربي: "Client-side. الخدمة المُستدعية (PDF Processing) هي التي تحمي نفسها."',
         'In English: "Client-side. The calling service (PDF Processing) is the one protecting itself."')
    ]

    for qa in qa_cb:
        add_heading(doc, qa[0], 3)
        add_paragraph_with_direction(doc, qa[1], is_quote=True)
        add_paragraph_with_direction(doc, qa[2], is_quote=True)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # Requirement 2: Strategy Pattern
    # ═══════════════════════════════════════════════════════════════

    add_heading(doc, '═' * 30, 1)
    add_heading(doc, 'المطلب الثاني: Strategy Pattern + Reflection', 1)
    add_heading(doc, '═' * 30, 1)

    add_heading(doc, 'ماذا تقول للمشرف (المقدمة):', 2)
    add_paragraph_with_direction(doc,
        'بالعربي: "هذا النمط يسمح بتبديل الخوارزميات في وقت التشغيل. مثلاً المستخدم يختار صيغة التقرير (Excel أو PDF أو CSV)، والنظام ينفذ الاستراتيجية المناسبة بدون if-else طويلة."',
        is_quote=True)
    add_paragraph_with_direction(doc,
        'In English: "This pattern allows swapping algorithms at runtime. For example, the user chooses the report format (Excel, PDF, or CSV), and the system executes the appropriate strategy without long if-else chains."',
        is_quote=True)

    add_heading(doc, 'ماذا تقول (Reflection):', 2)
    add_paragraph_with_direction(doc,
        'بالعربي: "استخدمت Reflection لتحميل الـ classes ديناميكياً من أسماء نصية. يعني أكتب \'csv\' كـ string، والنظام يحمّل CSVReportStrategy تلقائياً."',
        is_quote=True)
    add_paragraph_with_direction(doc,
        'In English: "I used Reflection to load classes dynamically from string names. Meaning I write \'csv\' as a string, and the system loads CSVReportStrategy automatically."',
        is_quote=True)

    add_heading(doc, 'الملفات ومواقع الكود', 2)

    add_heading(doc, 'الملف الرئيسي (الواجهة المجردة):', 3)
    add_paragraph_with_direction(doc,
        'المسار الكامل: c:\\Users\\HP\\smart-attendance-system\\reporting-service\\strategies\\report_strategy.py')

    add_heading(doc, 'تعريف Abstract Class (سطر 16-47):', 3)
    code_strategy = '''class ReportFormatStrategy(ABC):
    """Abstract base class for all report strategies"""

    @abstractmethod
    def generate_student_report(self,
                               student_data: Dict[str, Any],
                               attendance_records: List[Dict[str, Any]],
                               course_data: Dict[str, Any]) -> str:
        """Generate attendance report for one student"""
        pass

    @abstractmethod
    def generate_course_report(self,
                              course_data: Dict[str, Any],
                              lectures_data: List[Dict[str, Any]],
                              students_data: List[Dict[str, Any]],
                              attendance_matrix: Dict[tuple, str]) -> str:
        """Generate attendance report for entire course"""
        pass'''
    add_paragraph_with_direction(doc, code_strategy, is_code=True)

    add_paragraph_with_direction(doc, 'قل للمشرف:', is_quote=False)
    add_paragraph_with_direction(doc, 'بالعربي: "هذه الواجهة المجردة. كل استراتيجية جديدة لازم تنفّذ هذه الدوال."', is_quote=True)
    add_paragraph_with_direction(doc, 'In English: "This is the abstract interface. Every new strategy must implement these methods."', is_quote=True)

    add_heading(doc, 'المصنع مع Reflection:', 3)
    add_paragraph_with_direction(doc,
        'المسار الكامل: c:\\Users\\HP\\smart-attendance-system\\reporting-service\\strategy_factory.py')

    add_heading(doc, 'دالة create_strategy مع Reflection (سطر 156-189):', 3)
    code_factory = '''def create_strategy(self, format_name: Optional[str] = None) -> ReportFormatStrategy:
    """Create strategy using Reflection"""

    # Step 1: Determine format name
    if format_name is None:
        format_name = self.get_default_format()

    # Step 2: Convert name to class name
    class_name = self._format_name_to_class_name(format_name)
    # "csv" -> "CSVReportStrategy"

    # Step 3: Reflection - Load module dynamically
    strategies_module = importlib.import_module('strategies')

    # Step 4: Reflection - Get class from string name
    StrategyClass = getattr(strategies_module, class_name)

    # Step 5: Create instance
    return StrategyClass()'''
    add_paragraph_with_direction(doc, code_factory, is_code=True)

    add_paragraph_with_direction(doc, 'قل للمشرف:', is_quote=False)
    add_paragraph_with_direction(doc, 'بالعربي: "هنا السحر! importlib.import_module يحمّل الـ module ديناميكياً، و getattr يجيب الـ class من اسم نصي. هذا Reflection."', is_quote=True)
    add_paragraph_with_direction(doc, 'In English: "Here\'s the magic! importlib.import_module loads the module dynamically, and getattr gets the class from a string name. This is Reflection."', is_quote=True)

    add_heading(doc, 'كيف تُظهر للمشرف (خطوة بخطوة):', 2)

    demo_steps_strategy = [
        ('الخطوة 1: افتح Dashboard',
         'قل: "سأفتح صفحة Design Patterns ثم Strategy Pattern"',
         'Say: "I\'ll open the Design Patterns page then Strategy Pattern"',
         'http://localhost:3000 → Design Patterns → Strategy Pattern tab'),
        ('الخطوة 2: أظهر الصيغ المتاحة',
         'قل: "هنا 3 صيغ: Excel, PDF, CSV. كل واحدة استراتيجية منفصلة."',
         'Say: "Here are 3 formats: Excel, PDF, CSV. Each one is a separate strategy."',
         ''),
        ('الخطوة 3: ولّد تقرير Excel',
         'قل: "سأختار Excel وأضغط Generate"',
         'Say: "I\'ll choose Excel and click Generate"',
         'اختر Excel - اضغط Generate Report - أشر على الملف المحمّل'),
        ('الخطوة 4: ولّد تقرير PDF',
         'قل: "نفس الكود، بس غيّرت الاستراتيجية"',
         'Say: "Same code, just changed the strategy"',
         'اختر PDF - اضغط Generate Report')
    ]

    for step in demo_steps_strategy:
        add_heading(doc, step[0], 3)
        add_paragraph_with_direction(doc, step[1], is_quote=True)
        add_paragraph_with_direction(doc, step[2], is_quote=True)
        if step[3]:
            add_paragraph_with_direction(doc, step[3])

    add_heading(doc, 'أسئلة متوقعة:', 2)

    qa_strategy = [
        ('س: ما فائدة Strategy Pattern؟',
         'بالعربي: "Open/Closed Principle. أضيف صيغة جديدة بدون تعديل الكود الموجود."',
         'In English: "Open/Closed Principle. I add a new format without modifying existing code."'),
        ('س: ما فائدة Reflection؟',
         'بالعربي: "لا حاجة لـ if-else. أعطيه اسم نصي، يجيب الـ class تلقائياً."',
         'In English: "No need for if-else. I give it a string name, it gets the class automatically."'),
        ('س: كيف أضيف صيغة JSON؟',
         'بالعربي: "3 خطوات: 1) ملف json_strategy.py 2) class JSONReportStrategy 3) أضيف \'json\' في config.yml. انتهى!"',
         'In English: "3 steps: 1) json_strategy.py file 2) JSONReportStrategy class 3) add \'json\' in config.yml. Done!"')
    ]

    for qa in qa_strategy:
        add_heading(doc, qa[0], 3)
        add_paragraph_with_direction(doc, qa[1], is_quote=True)
        add_paragraph_with_direction(doc, qa[2], is_quote=True)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # Requirement 3: Choreography Pattern
    # ═══════════════════════════════════════════════════════════════

    add_heading(doc, '═' * 30, 1)
    add_heading(doc, 'المطلب الثالث: Choreography Pattern (RabbitMQ)', 1)
    add_heading(doc, '═' * 30, 1)

    add_heading(doc, 'ماذا تقول للمشرف (المقدمة):', 2)
    add_paragraph_with_direction(doc,
        'بالعربي: "هذا نمط تواصل بين الخدمات. بدل ما خدمة تتصل بخدمة مباشرة، تنشر حدث (Event) في message broker، والخدمات الأخرى تستهلك هذا الحدث بشكل مستقل. لا يوجد منسق مركزي."',
        is_quote=True)
    add_paragraph_with_direction(doc,
        'In English: "This is a communication pattern between services. Instead of one service calling another directly, it publishes an event to a message broker, and other services consume this event independently. There\'s no central coordinator."',
        is_quote=True)

    add_heading(doc, 'ماذا تقول (الفرق عن Orchestration):', 2)
    add_paragraph_with_direction(doc,
        'بالعربي: "في Orchestration، خدمة واحدة تتحكم بالكل مثل قائد الأوركسترا. في Choreography، كل خدمة ترقص لحالها عند سماع الموسيقى (Event)."',
        is_quote=True)
    add_paragraph_with_direction(doc,
        'In English: "In Orchestration, one service controls everything like an orchestra conductor. In Choreography, each service dances on its own when it hears the music (Event)."',
        is_quote=True)

    add_heading(doc, 'الملفات ومواقع الكود', 2)

    add_heading(doc, 'ملف RabbitMQ Client:', 3)
    add_paragraph_with_direction(doc,
        'المسار الكامل: c:\\Users\\HP\\smart-attendance-system\\common\\rabbitmq_client.py')

    add_heading(doc, 'تعريف الـ Client (سطر 20-46):', 3)
    code_rmq = '''class RabbitMQClient:
    """Simple RabbitMQ client for publishing and consuming messages"""

    QUEUE_NAME = 'attendance_events'

    def __init__(self, host: str = None):
        self.host = host or os.getenv('RABBITMQ_HOST', 'localhost')
        self.connection = None
        self.channel = None
        self._connect()

    def _connect(self):
        """Connect to RabbitMQ server"""
        self.connection = pika.BlockingConnection(
            pika.ConnectionParameters(host=self.host)
        )
        self.channel = self.connection.channel()
        self.channel.queue_declare(queue=self.QUEUE_NAME, durable=True)'''
    add_paragraph_with_direction(doc, code_rmq, is_code=True)

    add_paragraph_with_direction(doc, 'قل للمشرف:', is_quote=False)
    add_paragraph_with_direction(doc, 'بالعربي: "هذا الـ client للاتصال بـ RabbitMQ. الـ queue اسمها attendance_events."', is_quote=True)
    add_paragraph_with_direction(doc, 'In English: "This is the client for connecting to RabbitMQ. The queue is named attendance_events."', is_quote=True)

    add_heading(doc, 'دالة النشر publish (سطر 48-65):', 3)
    code_publish = '''def publish(self, message: dict):
    """Publish message (Producer)"""
    self.channel.basic_publish(
        exchange='',
        routing_key=self.QUEUE_NAME,
        body=json.dumps(message),
        properties=pika.BasicProperties(
            delivery_mode=2,  # Make message persistent
            content_type='application/json'
        )
    )'''
    add_paragraph_with_direction(doc, code_publish, is_code=True)

    add_paragraph_with_direction(doc, 'قل للمشرف:', is_quote=False)
    add_paragraph_with_direction(doc, 'بالعربي: "هذه دالة النشر. delivery_mode=2 يعني الرسالة تُحفظ على القرص."', is_quote=True)
    add_paragraph_with_direction(doc, 'In English: "This is the publish function. delivery_mode=2 means the message is saved to disk."', is_quote=True)

    add_heading(doc, 'نشر الحدث بعد تسجيل الحضور (Attendance Service سطر 91-105):', 3)
    add_paragraph_with_direction(doc,
        'المسار الكامل: c:\\Users\\HP\\smart-attendance-system\\attendance-service\\app.py')
    code_produce = '''if success:
    # Choreography: Publish event to RabbitMQ after successful recording
    try:
        rabbitmq.publish({
            'event': 'attendance_recorded',
            'student_id': student_id,
            'course_id': course_id,
            'date': date,
            'status': status
        })
        print(f"[RabbitMQ] Published attendance event for student {student_id}")
    except Exception as rmq_err:
        print(f"[RabbitMQ] Failed to publish event: {rmq_err}")
        raise'''
    add_paragraph_with_direction(doc, code_produce, is_code=True)

    add_paragraph_with_direction(doc, 'قل للمشرف:', is_quote=False)
    add_paragraph_with_direction(doc, 'بالعربي: "بعد تسجيل الحضور بنجاح، ننشر حدث إلى RabbitMQ. الخدمات الأخرى تستهلكه."', is_quote=True)
    add_paragraph_with_direction(doc, 'In English: "After successfully recording attendance, we publish an event to RabbitMQ. Other services consume it."', is_quote=True)

    add_heading(doc, 'كيف تُظهر للمشرف (خطوة بخطوة):', 2)

    demo_steps_choreo = [
        ('الخطوة 1: افتح RabbitMQ Management',
         'قل: "أولاً أظهر لك RabbitMQ"',
         'Say: "First I\'ll show you RabbitMQ"',
         'http://localhost:15672 - Username: guest - Password: guest'),
        ('الخطوة 2: أظهر الـ Queue',
         'قل: "هذه الـ queue اللي نستخدمها: attendance_events"',
         'Say: "This is the queue we use: attendance_events"',
         'اذهب إلى Queues → attendance_events'),
        ('الخطوة 3: افتح Dashboard Choreography tab',
         'قل: "الآن نرجع للـ Dashboard"',
         'Say: "Now let\'s go back to the Dashboard"',
         'http://localhost:3000 → Design Patterns → Choreography tab'),
        ('الخطوة 4: سجّل حضور (من OMR)',
         'قل: "سأسجّل حضور عبر OMR Processing"',
         'Say: "I\'ll record attendance via OMR Processing"',
         'اذهب إلى OMR Processing - ارفع bubble sheet - أو اضغط "Publish Test Event"'),
        ('الخطوة 5: أظهر الحدث في Consumed Events',
         'قل: "لاحظ: الحدث وصل Course Service تلقائياً"',
         'Say: "Notice: the event reached Course Service automatically"',
         'أشر على قائمة Events')
    ]

    for step in demo_steps_choreo:
        add_heading(doc, step[0], 3)
        add_paragraph_with_direction(doc, step[1], is_quote=True)
        add_paragraph_with_direction(doc, step[2], is_quote=True)
        add_paragraph_with_direction(doc, step[3])

    add_heading(doc, 'أسئلة متوقعة:', 2)

    qa_choreo = [
        ('س: ما الفرق بين Choreography و Orchestration؟',
         'بالعربي: "Orchestration: منسق مركزي يتحكم بكل شي. Choreography: كل خدمة تتصرف بشكل مستقل عند استلام Event."',
         'In English: "Orchestration: central coordinator controls everything. Choreography: each service acts independently when receiving an Event."'),
        ('س: لماذا RabbitMQ؟',
         'بالعربي: "Message Broker موثوق، يدعم persistence، وسهل الاستخدام مع Python."',
         'In English: "Reliable Message Broker, supports persistence, and easy to use with Python."'),
        ('س: ما معنى durable و delivery_mode=2؟',
         'بالعربي: "الرسائل تُحفظ على القرص. لو RabbitMQ أعاد التشغيل، الرسائل لا تضيع."',
         'In English: "Messages are saved to disk. If RabbitMQ restarts, messages are not lost."')
    ]

    for qa in qa_choreo:
        add_heading(doc, qa[0], 3)
        add_paragraph_with_direction(doc, qa[1], is_quote=True)
        add_paragraph_with_direction(doc, qa[2], is_quote=True)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # Requirement 4: JWT Authentication
    # ═══════════════════════════════════════════════════════════════

    add_heading(doc, '═' * 30, 1)
    add_heading(doc, 'المطلب الرابع: JWT Authentication', 1)
    add_heading(doc, '═' * 30, 1)

    add_heading(doc, 'ماذا تقول للمشرف (المقدمة):', 2)
    add_paragraph_with_direction(doc,
        'بالعربي: "JWT هو JSON Web Token. طريقة للـ authentication بدون session. المستخدم يسجل دخول، يحصل على token، ويرسله مع كل طلب. الـ token يحتوي معلومات المستخدم ومشفر بتوقيع."',
        is_quote=True)
    add_paragraph_with_direction(doc,
        'In English: "JWT is JSON Web Token. It\'s an authentication method without sessions. The user logs in, gets a token, and sends it with every request. The token contains user information and is secured with a signature."',
        is_quote=True)

    add_heading(doc, 'ماذا تقول (لماذا JWT):', 2)
    add_paragraph_with_direction(doc,
        'بالعربي: "في Microservices، الـ session صعب لأن كل خدمة منفصلة. JWT يحل المشكلة: التوكن يحمل معلومات المستخدم، وأي خدمة تقدر تتحقق منه."',
        is_quote=True)
    add_paragraph_with_direction(doc,
        'In English: "In Microservices, sessions are difficult because each service is separate. JWT solves the problem: the token carries user information, and any service can verify it."',
        is_quote=True)

    add_heading(doc, 'الملفات ومواقع الكود', 2)

    add_heading(doc, 'ملف Auth Service:', 3)
    add_paragraph_with_direction(doc,
        'المسار الكامل: c:\\Users\\HP\\smart-attendance-system\\auth-service\\app.py')

    add_heading(doc, 'المفتاح السري (سطر 25):', 3)
    code_secret = '''SECRET_KEY = os.getenv('JWT_SECRET_KEY', 'your-secret-key-change-in-production-2024')'''
    add_paragraph_with_direction(doc, code_secret, is_code=True)

    add_paragraph_with_direction(doc, 'قل للمشرف:', is_quote=False)
    add_paragraph_with_direction(doc, 'بالعربي: "هذا المفتاح السري للتوقيع. لازم يكون سري ونفسه في كل الخدمات."', is_quote=True)
    add_paragraph_with_direction(doc, 'In English: "This is the secret key for signing. It must be secret and the same across all services."', is_quote=True)

    add_heading(doc, 'دالة تسجيل الدخول login (سطر 146-187):', 3)
    code_login = '''@app.route('/api/auth/login', methods=['POST'])
def login():
    data = request.get_json()
    username = data.get('username')
    password = data.get('password')

    # Verify user
    user = db.fetch_one(
        "SELECT * FROM users WHERE username = ? AND is_active = 1",
        (username,)
    )

    if not user or user['password_hash'] != hash_password(password):
        return jsonify({'error': 'Invalid credentials'}), 401

    # Create JWT Token
    expiration = datetime.datetime.utcnow() + datetime.timedelta(hours=24)
    token = jwt.encode({
        'user_id': user['id'],
        'username': user['username'],
        'role': user['role'],
        'exp': expiration
    }, SECRET_KEY, algorithm='HS256')

    return jsonify({
        'token': token,
        'username': user['username'],
        'role': user['role']
    }), 200'''
    add_paragraph_with_direction(doc, code_login, is_code=True)

    add_paragraph_with_direction(doc, 'قل للمشرف:', is_quote=False)
    add_paragraph_with_direction(doc, 'بالعربي: "هنا نتحقق من الـ credentials، ثم ننشئ JWT token يحتوي user_id, username, role, exp."', is_quote=True)
    add_paragraph_with_direction(doc, 'In English: "Here we verify the credentials, then create a JWT token containing user_id, username, role, exp."', is_quote=True)

    add_heading(doc, 'دالة التحقق من التوكن في API Gateway (سطر 35-61):', 3)
    add_paragraph_with_direction(doc,
        'المسار الكامل: c:\\Users\\HP\\smart-attendance-system\\api-gateway\\app.py')
    code_validate = '''def validate_token():
    """Validate JWT token from request headers"""
    token = request.headers.get('Authorization')

    if not token:
        return None, {'error': 'Token required'}, 401

    try:
        # Remove "Bearer " prefix
        if token.startswith('Bearer '):
            token = token[7:]

        # Decode token and verify signature
        payload = jwt.decode(token, SECRET_KEY, algorithms=['HS256'])
        return payload, None, None

    except jwt.ExpiredSignatureError:
        return None, {'error': 'Token expired'}, 401
    except jwt.InvalidTokenError:
        return None, {'error': 'Invalid token'}, 401'''
    add_paragraph_with_direction(doc, code_validate, is_code=True)

    add_paragraph_with_direction(doc, 'قل للمشرف:', is_quote=False)
    add_paragraph_with_direction(doc, 'بالعربي: "الـ Gateway يتحقق من كل توكن. إذا صحيح، يمرر الطلب. إذا لا، يرجع 401."', is_quote=True)
    add_paragraph_with_direction(doc, 'In English: "The Gateway validates every token. If valid, it forwards the request. If not, it returns 401."', is_quote=True)

    add_heading(doc, 'كيف تُظهر للمشرف (خطوة بخطوة):', 2)

    demo_steps_jwt = [
        ('الخطوة 1: افتح Dashboard JWT tab',
         'قل: "سأفتح صفحة JWT Authentication"',
         'Say: "I\'ll open the JWT Authentication page"',
         'http://localhost:3000 → Design Patterns → JWT Authentication tab'),
        ('الخطوة 2: أظهر حالة "No Token"',
         'قل: "الآن ما في توكن. لاحظ البطاقة الرمادية."',
         'Say: "Now there\'s no token. Notice the gray card."',
         ''),
        ('الخطوة 3: سجّل الدخول',
         'قل: "سأسجل دخول بـ admin / admin123"',
         'Say: "I\'ll log in with admin / admin123"',
         'أدخل admin / admin123 - اضغط Login - أشر على التوكن اللي ظهر'),
        ('الخطوة 4: اشرح محتويات التوكن',
         'قل: "لاحظ: user_id, username, role, expires. كل هذا في التوكن."',
         'Say: "Notice: user_id, username, role, expires. All this is in the token."',
         'أشر على البطاقة الخضراء'),
        ('الخطوة 5: اختبر مع توكن',
         'قل: "الآن أختبر endpoint محمي مع التوكن"',
         'Say: "Now I\'ll test a protected endpoint with the token"',
         'اضغط "Access WITH Token" - أشر على النجاح الأخضر'),
        ('الخطوة 6: اختبر بدون توكن',
         'قل: "الآن بدون توكن"',
         'Say: "Now without a token"',
         'اضغط "Access WITHOUT Token" - أشر على الرفض: "401 Unauthorized"')
    ]

    for step in demo_steps_jwt:
        add_heading(doc, step[0], 3)
        add_paragraph_with_direction(doc, step[1], is_quote=True)
        add_paragraph_with_direction(doc, step[2], is_quote=True)
        if step[3]:
            add_paragraph_with_direction(doc, step[3])

    add_heading(doc, 'أسئلة متوقعة:', 2)

    qa_jwt = [
        ('س: ما الخوارزمية المستخدمة؟',
         'بالعربي: "HS256، يعني HMAC-SHA256. خوارزمية توقيع متماثلة."',
         'In English: "HS256, meaning HMAC-SHA256. A symmetric signing algorithm."'),
        ('س: هل JWT تشفير؟',
         'بالعربي: "لا! JWT توقيع وليس تشفير. الـ Payload مرئية لأي شخص، لكن لا يمكن تعديلها بدون المفتاح."',
         'In English: "No! JWT is signing, not encryption. The Payload is visible to anyone, but cannot be modified without the key."'),
        ('س: ما فائدة exp؟',
         'بالعربي: "يحدد انتهاء صلاحية التوكن. هنا 24 ساعة."',
         'In English: "It sets the token expiration. Here it\'s 24 hours."'),
        ('س: لماذا Stateless أفضل للـ Microservices؟',
         'بالعربي: "لا حاجة لمشاركة session بين الخدمات. كل خدمة تتحقق من التوكن بنفسها."',
         'In English: "No need to share sessions between services. Each service verifies the token itself."')
    ]

    for qa in qa_jwt:
        add_heading(doc, qa[0], 3)
        add_paragraph_with_direction(doc, qa[1], is_quote=True)
        add_paragraph_with_direction(doc, qa[2], is_quote=True)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # Summary Table
    # ═══════════════════════════════════════════════════════════════

    add_heading(doc, 'ملخص سريع (للمراجعة قبل العرض)', 1)

    # Create summary table
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'

    # Headers
    headers = ['المطلب', 'النمط', 'الملف الرئيسي', 'ماذا يفعل']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True

    # Data rows
    data = [
        ['1', 'Circuit Breaker', 'common\\circuit_breaker.py', 'يحمي من الانهيار المتتالي'],
        ['2', 'Strategy + Reflection', 'reporting-service\\strategy_factory.py', 'يبدّل صيغ التقارير ديناميكياً'],
        ['3', 'Choreography', 'common\\rabbitmq_client.py', 'تواصل غير متزامن عبر Events'],
        ['4', 'JWT Auth', 'auth-service\\app.py', 'توثيق بدون session']
    ]

    for row_idx, row_data in enumerate(data):
        for col_idx, cell_data in enumerate(row_data):
            table.rows[row_idx + 1].cells[col_idx].text = cell_data

    doc.add_paragraph()

    add_heading(doc, 'الروابط السريعة', 2)
    links = [
        'Dashboard: http://localhost:3000',
        'API Gateway: http://localhost:5000',
        'RabbitMQ: http://localhost:15672 (guest/guest)',
        'Auth Service: http://localhost:5007'
    ]
    for link in links:
        add_paragraph_with_direction(doc, link)

    add_heading(doc, 'أوامر التشغيل', 2)
    commands = '''.\\START.ps1     # تشغيل
.\\STOP.ps1      # إيقاف
Invoke-RestMethod -Uri "http://localhost:5000/api/health"  # فحص'''
    add_paragraph_with_direction(doc, commands, is_code=True)

    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('بالتوفيق في العرض!')
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'Arial'
    set_rtl(p)

    # Save
    output_path = 'PRESENTATION_GUIDE.docx'
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path

if __name__ == '__main__':
    create_word_document()
